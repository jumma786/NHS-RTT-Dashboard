"""
Generates the project report as a Word document (report/NHS_RTT_Report.docx).

Structured to match the assignment brief sections:
  1. Dataset Selection (3.1)
  2. Dataset Evaluation and Documentation (3.2)
  3. End-Consumer Identification (3.3)
  4. Perspective and Analysis Definition (3.4)
  5. Dashboard Design and Development (3.5)
  + Key Findings, Conclusion, References
"""

from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

DATA_DIR = Path(__file__).resolve().parent.parent / "data" / "processed"
OUT_DIR = Path(__file__).resolve().parent.parent / "report"


def fmt(n):
    if isinstance(n, float):
        return f"{n:,.1f}"
    return f"{n:,}"


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def build_report():
    nat = pd.read_csv(DATA_DIR / "rtt_national.csv")
    icb = pd.read_csv(DATA_DIR / "rtt_icb.csv")

    trends = nat.groupby("Period").agg(
        total=("Total number of incomplete pathways", "sum"),
        within_18=("Total within 18 weeks", "sum"),
        w52=("Total 52 plus weeks", "sum"),
        w65=("Total 65 plus weeks", "sum"),
        w78=("Total 78 plus weeks", "sum"),
    ).reset_index()
    trends["pct_18"] = trends["within_18"] / trends["total"] * 100
    trends["pct_52"] = trends["w52"] / trends["total"] * 100

    n_periods = nat["Period"].nunique()
    p_min, p_max = nat["Period"].min(), nat["Period"].max()
    first = trends[trends["Period"] == p_min].iloc[0]
    last = trends[trends["Period"] == p_max].iloc[0]
    peak = trends.loc[trends["total"].idxmax()]

    n_icbs = icb["ICB Code"].nunique()
    n_tfs = nat["Treatment Function Code"].nunique()
    n_rows_icb = len(icb)

    doc = Document()

    # ── Styles ──
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    for level in range(1, 4):
        doc.styles[f"Heading {level}"].font.color.rgb = RGBColor(0, 51, 102)

    # ── Title Page ──
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("NHS RTT Waiting Times Dashboard")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0, 51, 102)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "Analysis of Referral to Treatment Incomplete Pathways\n"
        "at Integrated Care Board Level"
    )
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"October 2022 – March 2026 ({n_periods} months)")
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "UK Data Dashboard Project\n"
        "Data Analytics Programme — Skill Versed"
    )
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(120, 120, 120)

    doc.add_page_break()

    # ── Table of Contents ──
    doc.add_heading("Table of Contents", level=1)
    doc.add_paragraph(
        "1. Introduction\n"
        "2. Dataset Selection\n"
        "3. Dataset Evaluation and Documentation\n"
        "    3.1 Description\n"
        "    3.2 Reasons for Selection\n"
        "    3.3 Key Strengths\n"
        "    3.4 Insights the Dataset Can Provide\n"
        "4. End-Consumer Identification\n"
        "    4.1 Potential End Users\n"
        "    4.2 Primary Target Audience\n"
        "    4.3 Justification\n"
        "5. Perspective and Analysis Definition\n"
        "    5.1 Alternative Perspectives\n"
        "    5.2 Chosen Perspective\n"
        "    5.3 Rationale\n"
        "6. Key Findings\n"
        "    6.1 National Overview\n"
        "    6.2 Waiting List Trends\n"
        "    6.3 Long-Wait Reduction\n"
        "    6.4 Specialty Analysis\n"
        "    6.5 ICB-Level Performance\n"
        "    6.6 Regional Variation and Improvement\n"
        "7. Dashboard Design and Development\n"
        "    7.1 Layout Decisions\n"
        "    7.2 Visualisation Choices\n"
        "    7.3 User Experience Considerations\n"
        "    7.4 Benefits of This Approach\n"
        "8. Conclusion\n"
        "9. References"
    )
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # SECTION 1: INTRODUCTION
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "Referral to Treatment (RTT) waiting times are one of the most important "
        "measures of NHS performance in England. They track how long patients wait "
        "from the point of referral by a GP or other healthcare professional to the "
        "start of their treatment. The NHS Constitution sets a standard that 92% of "
        "patients on incomplete (non-emergency) pathways should be waiting no more "
        "than 18 weeks."
    )
    doc.add_paragraph(
        f"This report analyses RTT incomplete pathway data published by NHS England "
        f"over {n_periods} months from October 2022 to March 2026. It examines trends at "
        "national, specialty, and Integrated Care Board (ICB) levels, with a focus on "
        "understanding how waiting list volumes and long-wait cohorts have changed "
        "over the period."
    )
    doc.add_paragraph(
        "The report is structured around the five project requirements set out in the "
        "assignment brief: dataset selection, dataset evaluation, end-consumer "
        "identification, analytical perspective, and dashboard design. Key findings "
        "and a conclusion follow."
    )

    # ═══════════════════════════════════════════════════════════════════
    # SECTION 2: DATASET SELECTION (Brief 3.1)
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("2. Dataset Selection", level=1)
    doc.add_paragraph(
        "The dataset selected for this project is the NHS England Referral to Treatment "
        "(RTT) Waiting Times statistics, specifically the “Incomplete Pathways – "
        "Commissioner” monthly files. These are published by NHS England and are "
        "freely available from:"
    )
    doc.add_paragraph(
        "https://www.england.nhs.uk/statistics/statistical-work-areas/rtt-waiting-times/"
    )
    doc.add_paragraph(
        "This dataset falls within the “Public Services (NHS)” and "
        "“Healthcare (hospital performance)” domains recommended in the "
        "project brief. NHS England is listed as one of the recommended data sources. "
        "The data represents a broad UK domain with clear public and societal relevance "
        "— NHS waiting times are regularly reported in national media and directly "
        "affect millions of patients."
    )
    doc.add_paragraph(
        f"A total of {n_periods} monthly files were downloaded, covering October 2022 to "
        "March 2026. April 2023 was unavailable from the source at the time of "
        "collection. Each monthly file is an Excel workbook (.xlsx) containing "
        "data at National, Regional, ICB, and Sub-ICB levels, broken down by "
        "treatment function (medical specialty)."
    )

    # ═══════════════════════════════════════════════════════════════════
    # SECTION 3: DATASET EVALUATION AND DOCUMENTATION (Brief 3.2)
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("3. Dataset Evaluation and Documentation", level=1)

    doc.add_heading("3.1 Description", level=2)
    doc.add_paragraph(
        "The dataset contains monthly counts of patients on incomplete RTT pathways, "
        "broken down by week-band (how many weeks they have been waiting). "
        "Key fields include:"
    )
    doc.add_paragraph(
        f"Time period: {n_periods} months (October 2022 to March 2026).",
        style="List Bullet",
    )
    doc.add_paragraph(
        f"Geographic scope: All {n_icbs} Integrated Care Boards (ICBs) across England.",
        style="List Bullet",
    )
    doc.add_paragraph(
        f"Granularity: Monthly, per ICB, per treatment function ({n_tfs} categories "
        "including a Total row).",
        style="List Bullet",
    )
    doc.add_paragraph(
        f"Volume: {fmt(n_rows_icb)} ICB-level rows after processing.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Measures: 104 weekly wait-band columns (>0–1 through >103–104, "
        "plus “104 plus”), total incomplete pathways, total within 18 weeks, "
        "% within 18 weeks, median waiting time, 92nd percentile waiting time, and "
        "long-wait totals (52+, 65+, 78+ weeks).",
        style="List Bullet",
    )

    doc.add_heading("3.2 Reasons for Selection", level=2)
    doc.add_paragraph(
        "NHS waiting times were chosen for several reasons:"
    )
    doc.add_paragraph(
        "Relevance: RTT waiting times are the single most-watched NHS performance "
        "metric and are central to the government’s elective recovery programme.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Richness: The data supports analysis across time, geography (ICB), and "
        "clinical specialty — giving multiple dimensions for dashboard interactivity.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Public interest: Millions of patients are directly affected by waiting times, "
        "making this a dataset with genuine societal relevance.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Timeliness: The data is updated monthly, and the files used extend to "
        "March 2026 — the most recent available at the time of the project.",
        style="List Bullet",
    )

    doc.add_heading("3.3 Key Strengths", level=2)
    doc.add_paragraph(
        "Completeness: Zero null values or missing rows across all 41,328 ICB records. "
        "Every ICB reports every specialty in every month.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Consistency: The file structure is identical across all 41 months — same "
        "columns, same header layout, same ICB and treatment function codes.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Credibility: Published by NHS England, the official body responsible for "
        "commissioning healthcare in England. This is the same data used in "
        "parliamentary reporting.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Recency: Data extends to March 2026, providing a current picture of "
        "NHS performance.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Accuracy: Wait-band column sums exactly match published totals in every "
        "row — verified during the ETL process.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Known caveats (documented by NHS England):"
    )
    doc.add_paragraph(
        "October and November 2022: Missing data for Frimley Health NHS FT (RDU) "
        "and Manchester University NHS FT (R0A).", style="List Bullet"
    )
    doc.add_paragraph(
        "December 2022: Missing data for Manchester University NHS FT (R0A).",
        style="List Bullet",
    )
    doc.add_paragraph(
        "April 2023: No file was available from the source at the time of collection.",
        style="List Bullet",
    )

    doc.add_heading("3.4 Insights the Dataset Can Provide", level=2)
    doc.add_paragraph(
        "The dataset enables the following types of analysis:"
    )
    doc.add_paragraph(
        "National trend monitoring — how the total waiting list and key metrics "
        "(% within 18 weeks, median wait) have changed month by month.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Long-wait elimination tracking — progress against 52-week, 65-week, "
        "and 78-week targets.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "ICB benchmarking — comparing performance across all 42 ICBs to "
        "identify best and worst performers.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Specialty analysis — understanding which clinical specialties carry "
        "the largest waiting lists and where performance is weakest.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Year-on-year comparison — measuring improvement or deterioration "
        "against the same month in the prior year.",
        style="List Bullet",
    )

    # ═══════════════════════════════════════════════════════════════════
    # SECTION 4: END-CONSUMER IDENTIFICATION (Brief 3.3)
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("4. End-Consumer Identification", level=1)

    doc.add_heading("4.1 Potential End Users", level=2)
    doc.add_paragraph(
        "The following stakeholder groups may find this dashboard relevant:"
    )
    doc.add_paragraph(
        "ICB Performance and Planning Teams — responsible for managing local "
        "waiting lists, allocating capacity, and reporting to NHS England.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "NHS England National Programme Directors — who monitor system-wide "
        "progress against elective recovery targets and allocate funding.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Hospital Trust Operational Managers — who manage theatre schedules, "
        "outpatient capacity, and specialty-level demand.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Health Select Committee and Policy Advisors — who scrutinise NHS "
        "performance and advise on policy interventions.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Healthcare Journalists and Analysts — who report on NHS performance "
        "to the public and need accessible, accurate data summaries.",
        style="List Bullet",
    )

    doc.add_heading("4.2 Primary Target Audience", level=2)
    doc.add_paragraph(
        "The primary target audience for this dashboard is ICB Performance and "
        "Planning Teams. These are the operational decision-makers who directly "
        "manage waiting list reduction at the local level."
    )

    doc.add_heading("4.3 Justification", level=2)
    doc.add_paragraph(
        "ICB teams are the most relevant and valuable audience for this dashboard "
        "because they are the primary unit of accountability for elective waiting "
        "times since the 2022 ICB restructure. Specifically, they need to:"
    )
    doc.add_paragraph(
        "Benchmark their ICB’s performance against national averages and "
        "peer ICBs to understand their relative position.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Identify which specialties are driving long waits in their area, so they "
        "can target capacity investment.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Track month-on-month and year-on-year trends to assess whether local "
        "interventions (e.g. insourcing, waiting list initiatives) are working.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Report upward to NHS England with evidence-based summaries of progress "
        "against national targets.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "The dashboard is designed to support all four of these needs through "
        "filterable data, comparative charts, pre-computed KPIs, and ranked "
        "ICB tables. Other audiences (journalists, policy advisors) benefit from "
        "the same dashboard but are secondary users."
    )

    # ═══════════════════════════════════════════════════════════════════
    # SECTION 5: PERSPECTIVE AND ANALYSIS DEFINITION (Brief 3.4)
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("5. Perspective and Analysis Definition", level=1)

    doc.add_heading("5.1 Alternative Perspectives", level=2)
    doc.add_paragraph(
        "The RTT dataset can be analysed from several different angles:"
    )
    doc.add_paragraph(
        "Capacity Planning — focusing on total volumes and demand forecasting "
        "to inform staffing and theatre allocation. This perspective would prioritise "
        "absolute numbers and growth rates.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Patient Experience — focusing on median and 92nd percentile wait times "
        "to understand what patients actually experience. This perspective would "
        "prioritise distribution analysis.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Equity and Variation — focusing on geographic disparities between ICBs "
        "to highlight inequalities in access to care. This perspective would "
        "prioritise ICB-level comparisons and spread.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Elective Recovery Progress — focusing on whether the NHS is meeting "
        "its targets to reduce waiting lists and eliminate the longest waits "
        "(52+, 65+, 78+ weeks). This perspective combines trend analysis with "
        "target tracking.",
        style="List Bullet",
    )

    doc.add_heading("5.2 Chosen Perspective", level=2)
    doc.add_paragraph(
        "This dashboard focuses on Elective Recovery Progress. The central "
        "question is:"
    )
    p = doc.add_paragraph()
    run = p.add_run(
        "“How effectively are ICBs reducing their waiting lists and eliminating "
        "long waits since the post-COVID recovery programme began?”"
    )
    run.italic = True

    doc.add_heading("5.3 Rationale", level=2)
    doc.add_paragraph(
        "This perspective was chosen because it directly aligns with the NHS’s "
        "own stated priorities. The 2023 Elective Recovery Plan set explicit targets "
        "to eliminate 78-week waits by April 2023, 65-week waits by March 2024, and "
        "to make significant progress on 52-week waits. This perspective matters "
        "because:"
    )
    doc.add_paragraph(
        "It is the lens through which NHS England itself evaluates ICB performance.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "It serves the primary audience (ICB planning teams) by providing "
        "actionable insight — not just “where are we now?” but "
        "“are we improving, and where should we focus next?”",
        style="List Bullet",
    )
    doc.add_paragraph(
        "It informs decisions about resource allocation, specialty-level "
        "investment, and waiting list initiatives.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "The year-on-year comparison sheet, the best/worst ICB rankings, and the "
        "long-wait trend charts all serve this perspective directly."
    )

    # ═══════════════════════════════════════════════════════════════════
    # SECTION 6: KEY FINDINGS
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("6. Key Findings", level=1)

    # 6.1 National Overview
    doc.add_heading("6.1 National Overview", level=2)
    doc.add_paragraph(
        f"As of March 2026, a total of {fmt(int(last['total']))} patients were on "
        f"incomplete RTT pathways nationally. This represents a "
        f"{(last['total'] - first['total']) / first['total'] * 100:+.1f}% change "
        f"from {fmt(int(first['total']))} in October 2022. The waiting list peaked "
        f"at {fmt(int(peak['total']))} in {peak['Period']} before declining gradually."
    )
    add_table(doc,
        ["Metric", "Oct 2022", "Mar 2026", "Change"],
        [
            ["Total incomplete pathways", fmt(int(first["total"])),
             fmt(int(last["total"])),
             f"{(last['total'] - first['total']) / first['total'] * 100:+.1f}%"],
            ["% within 18 weeks", f"{first['pct_18']:.1f}%",
             f"{last['pct_18']:.1f}%",
             f"{last['pct_18'] - first['pct_18']:+.1f}pp"],
            ["52+ week waiters", fmt(int(first["w52"])), fmt(int(last["w52"])),
             f"{(last['w52'] - first['w52']) / first['w52'] * 100:+.1f}%"],
            ["65+ week waiters", fmt(int(first["w65"])), fmt(int(last["w65"])),
             f"{(last['w65'] - first['w65']) / first['w65'] * 100:+.1f}%"],
            ["78+ week waiters", fmt(int(first["w78"])), fmt(int(last["w78"])),
             f"{(last['w78'] - first['w78']) / first['w78'] * 100:+.1f}%"],
        ],
        col_widths=[5, 3.5, 3.5, 3],
    )
    doc.add_paragraph("Table 1: National headline RTT metrics, October 2022 vs March 2026.")

    # 6.2 Waiting List Trends
    doc.add_heading("6.2 Waiting List Size Trends", level=2)
    doc.add_paragraph(
        f"The total waiting list grew from {fmt(int(first['total']))} in October 2022 "
        f"to a peak of {fmt(int(peak['total']))} in August 2023, driven by continued "
        f"post-pandemic demand outstripping capacity. From this peak, the list declined "
        f"steadily to {fmt(int(last['total']))} in March 2026 — a reduction of "
        f"{fmt(int(peak['total'] - last['total']))} "
        f"({(peak['total'] - last['total']) / peak['total'] * 100:.1f}%) from peak."
    )
    doc.add_paragraph(
        f"The proportion of patients seen within 18 weeks improved from "
        f"{first['pct_18']:.1f}% to {last['pct_18']:.1f}% over the period, though "
        f"this remains well below the 92% constitutional standard."
    )

    key_periods = ["2022-10", "2023-02", "2023-08", "2024-03",
                   "2024-09", "2025-03", "2025-09", "2026-03"]
    rows = []
    for per in key_periods:
        r = trends[trends["Period"] == per].iloc[0]
        rows.append([r["Period"], fmt(int(r["total"])), f"{r['pct_18']:.1f}%",
                     fmt(int(r["w52"])), f"{r['pct_52']:.2f}%"])
    add_table(doc,
        ["Period", "Total Pathways", "% Within 18w", "52+ Waiters", "% 52+ Weeks"],
        rows, col_widths=[2.5, 3.5, 3, 3.5, 3],
    )
    doc.add_paragraph("Table 2: Waiting list trajectory at selected time points.")

    # 6.3 Long Waits
    doc.add_heading("6.3 Long-Wait Reduction", level=2)
    doc.add_paragraph(
        "The most significant achievement over the analysis period has been the "
        "dramatic reduction in the longest-waiting cohorts:"
    )
    doc.add_paragraph(
        f"78+ week waiters fell by "
        f"{(first['w78'] - last['w78']) / first['w78'] * 100:.1f}%, "
        f"from {fmt(int(first['w78']))} to {fmt(int(last['w78']))}.",
        style="List Bullet",
    )
    doc.add_paragraph(
        f"65+ week waiters fell by "
        f"{(first['w65'] - last['w65']) / first['w65'] * 100:.1f}%, "
        f"from {fmt(int(first['w65']))} to {fmt(int(last['w65']))}.",
        style="List Bullet",
    )
    doc.add_paragraph(
        f"52+ week waiters fell by "
        f"{(first['w52'] - last['w52']) / first['w52'] * 100:.1f}%, "
        f"from {fmt(int(first['w52']))} to {fmt(int(last['w52']))}.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "This reflects the NHS’s stated priority of eliminating the longest "
        "waits first, with successive targets to eliminate 104-week, 78-week, and "
        "65-week waits. The data confirms substantial progress against all three."
    )

    # 6.4 Specialty Analysis
    doc.add_heading("6.4 Specialty Analysis", level=2)
    latest_nat = nat[nat["Period"] == p_max].copy()
    latest_nat["pct_18"] = (latest_nat["Total within 18 weeks"]
                            / latest_nat["Total number of incomplete pathways"] * 100)
    latest_nat = latest_nat[latest_nat["Treatment Function Code"] != "C_999"]

    doc.add_paragraph(
        "Performance varies considerably across treatment functions. As of March 2026:"
    )

    worst5 = latest_nat.nsmallest(5, "pct_18")
    spec_rows = [[r["Treatment Function"],
                   fmt(int(r["Total number of incomplete pathways"])),
                   f"{r['pct_18']:.1f}%"] for _, r in worst5.iterrows()]
    add_table(doc, ["Specialty", "Total Pathways", "% Within 18w"],
              spec_rows, col_widths=[7, 3.5, 3])
    doc.add_paragraph("Table 3: Worst-performing specialties (March 2026).")

    best5 = latest_nat.nlargest(5, "pct_18")
    spec_rows2 = [[r["Treatment Function"],
                    fmt(int(r["Total number of incomplete pathways"])),
                    f"{r['pct_18']:.1f}%"] for _, r in best5.iterrows()]
    add_table(doc, ["Specialty", "Total Pathways", "% Within 18w"],
              spec_rows2, col_widths=[7, 3.5, 3])
    doc.add_paragraph("Table 4: Best-performing specialties (March 2026).")

    doc.add_paragraph(
        "Oral Surgery and Plastic Surgery have the lowest proportion within 18 weeks, "
        "while Elderly Medicine and Mental Health services perform best. High-volume "
        "specialties such as Trauma & Orthopaedics and ENT remain under significant "
        "pressure."
    )

    # 6.5 ICB Performance
    doc.add_heading("6.5 ICB-Level Performance", level=2)
    latest_icb = icb[icb["Period"] == p_max].copy()
    icb_agg = latest_icb.groupby(["ICB Code", "ICB Name"]).agg(
        total=("Total number of incomplete pathways", "sum"),
        within_18=("Total within 18 weeks", "sum"),
        w52=("Total 52 plus weeks", "sum"),
    ).reset_index()
    icb_agg["pct_18"] = icb_agg["within_18"] / icb_agg["total"] * 100

    doc.add_paragraph(
        f"Across the 42 ICBs, performance in March 2026 ranged from "
        f"{icb_agg['pct_18'].min():.1f}% to {icb_agg['pct_18'].max():.1f}% within "
        f"18 weeks — a spread of "
        f"{icb_agg['pct_18'].max() - icb_agg['pct_18'].min():.1f} percentage points, "
        f"indicating substantial geographic inequality."
    )

    top5 = icb_agg.nlargest(5, "pct_18")
    icb_rows = [[r["ICB Code"], r["ICB Name"][:50], fmt(int(r["total"])),
                  f"{r['pct_18']:.1f}%"] for _, r in top5.iterrows()]
    add_table(doc, ["Code", "ICB Name", "Total Pathways", "% Within 18w"],
              icb_rows, col_widths=[1.5, 7, 3, 3])
    doc.add_paragraph("Table 5: Top 5 ICBs by % within 18 weeks (March 2026).")

    bot5 = icb_agg.nsmallest(5, "pct_18")
    icb_rows2 = [[r["ICB Code"], r["ICB Name"][:50], fmt(int(r["total"])),
                   f"{r['pct_18']:.1f}%"] for _, r in bot5.iterrows()]
    add_table(doc, ["Code", "ICB Name", "Total Pathways", "% Within 18w"],
              icb_rows2, col_widths=[1.5, 7, 3, 3])
    doc.add_paragraph("Table 6: Bottom 5 ICBs by % within 18 weeks (March 2026).")

    # 6.6 Regional Variation
    doc.add_heading("6.6 Regional Variation and Improvement", level=2)

    first_icb = icb[icb["Period"] == p_min].groupby("ICB Code").agg(
        t=("Total number of incomplete pathways", "sum"),
        w18=("Total within 18 weeks", "sum"),
    ).reset_index()
    first_icb["pct_start"] = first_icb["w18"] / first_icb["t"] * 100

    last_icb = icb[icb["Period"] == p_max].groupby("ICB Code").agg(
        t=("Total number of incomplete pathways", "sum"),
        w18=("Total within 18 weeks", "sum"),
    ).reset_index()
    last_icb["pct_end"] = last_icb["w18"] / last_icb["t"] * 100

    change = first_icb[["ICB Code", "pct_start"]].merge(
        last_icb[["ICB Code", "pct_end"]], on="ICB Code")
    change["improvement"] = change["pct_end"] - change["pct_start"]
    change = change.merge(icb_agg[["ICB Code", "ICB Name"]], on="ICB Code")

    doc.add_paragraph(
        f"The degree of improvement over the study period varies widely. "
        f"The most improved ICB gained {change['improvement'].max():.1f} percentage "
        f"points, while {(change['improvement'] < 0).sum()} ICBs actually saw their "
        f"performance decline."
    )

    imp_rows = [[r["ICB Code"], r["ICB Name"][:50], f"{r['pct_start']:.1f}%",
                  f"{r['pct_end']:.1f}%", f"+{r['improvement']:.1f}pp"]
                 for _, r in change.nlargest(5, "improvement").iterrows()]
    add_table(doc, ["Code", "ICB Name", "Oct 2022", "Mar 2026", "Change"],
              imp_rows, col_widths=[1.5, 6.5, 2, 2, 2])
    doc.add_paragraph("Table 7: Most improved ICBs (% within 18 weeks).")

    dec_rows = [[r["ICB Code"], r["ICB Name"][:50], f"{r['pct_start']:.1f}%",
                  f"{r['pct_end']:.1f}%", f"{r['improvement']:+.1f}pp"]
                 for _, r in change.nsmallest(5, "improvement").iterrows()]
    add_table(doc, ["Code", "ICB Name", "Oct 2022", "Mar 2026", "Change"],
              dec_rows, col_widths=[1.5, 6.5, 2, 2, 2])
    doc.add_paragraph("Table 8: ICBs with declining performance.")

    # ═══════════════════════════════════════════════════════════════════
    # SECTION 7: DASHBOARD DESIGN AND DEVELOPMENT (Brief 3.5)
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("7. Dashboard Design and Development", level=1)

    doc.add_heading("7.1 Layout Decisions", level=2)
    doc.add_paragraph(
        "The dashboard follows a top-down information hierarchy designed for quick "
        "scanning. The top banner uses NHS Identity colours (dark blue #003087) to "
        "establish credibility and context. Below it, four KPI cards provide an "
        "at-a-glance summary of the most critical metrics: total pathways, % within "
        "18 weeks, median wait, and 52+ week waiters. Each card shows the current "
        "value and the month-on-month change, so users can immediately see the "
        "direction of travel without scrolling."
    )
    doc.add_paragraph(
        "Charts are arranged in a two-column grid (left and right) with consistent "
        "sizing, followed by a full-width trend chart and an embedded table at the "
        "bottom. This layout ensures the dashboard reads naturally from top to bottom "
        "and left to right, with the most important information (KPIs) appearing first."
    )
    doc.add_paragraph(
        "The workbook contains eight sheets in a deliberate order: Dashboard (summary), "
        "Data (full dataset), National_Trend, ICB_Summary, Specialty_Top5, "
        "Top5_Bottom5, YoY_Change, and Lookups. This moves from overview to detail, "
        "supporting both quick-glance and deep-dive use cases."
    )

    doc.add_heading("7.2 Visualisation Choices", level=2)
    doc.add_paragraph(
        "Five chart types were selected to provide variety while serving distinct "
        "analytical purposes:"
    )
    doc.add_paragraph(
        "Line chart (monthly pathways trend) — chosen because time-series data "
        "is best represented by lines, clearly showing the rise to August 2023 peak "
        "and subsequent decline.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Pie chart (specialty share) — shows the proportion of the waiting list "
        "held by each major specialty. Limited to 5 slices plus “Other” to "
        "avoid visual clutter.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Dual-line chart (52+ vs 65+ week waiters) — enables direct comparison "
        "of two long-wait cohorts on the same axis, highlighting the faster reduction "
        "in 65+ week waits.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Horizontal bar chart (top 5 ICBs) — effective for comparing named "
        "categories, with ICB names readable on the y-axis without rotation.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Line chart (52+ week trend) — a focused view of the long-wait "
        "elimination target, coloured red to signal urgency and importance.",
        style="List Bullet",
    )

    doc.add_heading("7.3 User Experience Considerations", level=2)
    doc.add_paragraph(
        "Several design decisions were made to improve accessibility and usability "
        "for non-technical users:"
    )
    doc.add_paragraph(
        "The Data sheet is formatted as an Excel Table (“RTTData”) with "
        "auto-filters and frozen headers, allowing users to add Slicers (Insert > "
        "Slicer) for interactive filtering by ICB, specialty, or period without "
        "any advanced Excel knowledge.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Large numbers use K/M formatting (e.g. 7.0M, 92.6K) to improve readability "
        "at a glance, rather than displaying raw figures like 7,014,879.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Conditional formatting is applied throughout: data bars on the % within 18 "
        "weeks column, a red-amber-green colour scale on the ICB Summary sheet, and "
        "red/green highlighting on year-on-year change values to instantly convey "
        "good vs bad performance.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Grid lines are hidden on the Dashboard sheet and a light grey background "
        "(#F0F4F5) is applied to create a clean, report-like appearance that "
        "separates it visually from the data sheets.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "KPI values use dynamic Excel formulas referencing the National_Trend sheet, "
        "so they update automatically if new data is added.",
        style="List Bullet",
    )

    doc.add_heading("7.4 Benefits of This Approach", level=2)
    doc.add_paragraph(
        "The design is effective for the target audience (ICB planning teams) because "
        "it supports three distinct use modes:"
    )
    doc.add_paragraph(
        "Quick executive summary — the four KPI cards and key insights section "
        "give a 30-second overview of current performance and direction of travel.",
        style="List Number",
    )
    doc.add_paragraph(
        "Visual narrative — the five charts tell the story of elective recovery "
        "progress across time, specialty, and geography.",
        style="List Number",
    )
    doc.add_paragraph(
        "Deep-dive analysis — the filterable Data sheet with 39,606 rows allows "
        "analysts to explore any ICB, specialty, or time period in detail.",
        style="List Number",
    )
    doc.add_paragraph(
        "By keeping everything in a single .xlsx file with no macros or external "
        "dependencies, the dashboard can be shared via email, opened on any machine "
        "with Excel, and does not require specialist software or training. This "
        "maximises accessibility for the target audience."
    )

    # ═══════════════════════════════════════════════════════════════════
    # SECTION 8: CONCLUSION
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("8. Conclusion", level=1)
    doc.add_paragraph(
        f"This analysis of {n_periods} months of NHS RTT data reveals a mixed picture. "
        f"While the overall waiting list remains large at {fmt(int(last['total']))} "
        f"patients, there has been meaningful progress in several key areas:"
    )
    doc.add_paragraph(
        f"The proportion of patients within 18 weeks improved from "
        f"{first['pct_18']:.1f}% to {last['pct_18']:.1f}%, a gain of "
        f"{last['pct_18'] - first['pct_18']:.1f} percentage points.",
        style="List Number",
    )
    doc.add_paragraph(
        f"The longest waits have been dramatically reduced: 78+ week waiters fell "
        f"by {(first['w78'] - last['w78']) / first['w78'] * 100:.0f}% and 52+ week "
        f"waiters by {(first['w52'] - last['w52']) / first['w52'] * 100:.0f}%.",
        style="List Number",
    )
    doc.add_paragraph(
        f"The waiting list peaked in August 2023 at {fmt(int(peak['total']))} and "
        f"has since declined by "
        f"{(peak['total'] - last['total']) / peak['total'] * 100:.1f}%.",
        style="List Number",
    )
    doc.add_paragraph(
        "However, significant challenges remain. The 92% within-18-weeks "
        f"constitutional standard appears distant at {last['pct_18']:.1f}%. "
        f"Geographic inequality persists, with a "
        f"{icb_agg['pct_18'].max() - icb_agg['pct_18'].min():.1f} percentage point "
        f"spread between the best and worst ICBs. Specialties such as Oral Surgery, "
        f"Plastic Surgery, and Trauma & Orthopaedics remain under particular pressure."
    )
    doc.add_paragraph(
        "The accompanying Excel dashboard enables stakeholders to explore these "
        "trends interactively, filtering by ICB, specialty, and time period to "
        "support local planning and performance management."
    )

    # ═══════════════════════════════════════════════════════════════════
    # SECTION 9: REFERENCES
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("9. References", level=1)
    doc.add_paragraph(
        "NHS England (2026). Consultant-led Referral to Treatment Waiting Times. "
        "Available at: https://www.england.nhs.uk/statistics/statistical-work-areas/"
        "rtt-waiting-times/ (Accessed: June 2026)."
    )
    doc.add_paragraph(
        "Department of Health and Social Care (2024). The NHS Constitution for "
        "England. Available at: https://www.gov.uk/government/publications/"
        "the-nhs-constitution-for-england"
    )
    doc.add_paragraph(
        "NHS England (2023). Delivery plan for tackling the COVID-19 backlog of "
        "elective care."
    )

    # ── Save ──
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    out_path = OUT_DIR / "NHS_RTT_Report.docx"
    doc.save(str(out_path))
    print(f"Report saved to {out_path}")


if __name__ == "__main__":
    build_report()
